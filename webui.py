import sys,os
now_dir = os.getcwd()
sys.path.append(now_dir) # Add the working directory to the path
import gradio as gr
from ultralytics import YOLO
from PIL import Image
import torch.utils.data.distributed
import torchvision.transforms as transforms
from torch.autograd import Variable
from tools.refer import Refer
import torch
import json
import argparse
import docx, docx2pdf
from docx.shared import Cm, Inches
import pythoncom

language = 'en'
# cn -> Chinese
# en -> English

parser = argparse.ArgumentParser()
parser.add_argument('-l', '--language', type = str, default = 'en', help = 'Gradio Output Language')
args = parser.parse_args()
language = args.language

Refer = Refer(language) # Proper nouns references

model_yolo = YOLO('weight/detection/yolov8.pt') # Detection model
model_cnn = torch.load('weight/classification/resnet18.pth') # Classification model

# Device to run
DEVICE = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")

# Tranformation of the image
transform_test = transforms.Compose([
    transforms.Resize((224, 224)),
    transforms.ToTensor(),
    transforms.Normalize([0.5, 0.5, 0.5], [0.5, 0.5, 0.5])
])

css = ".output_text { font-size: 18px; }"
classes = ('Noleopard', 'Leopard')
severity = ('levelzero', 'levelone', 'leveltwo', 'levelthree')
model_cnn.eval() # Set the model to evaluation mode
model_cnn.to(DEVICE)

def predict_image(name, img_path):
    doc_out = docx.Document()

    output_str = 'Patient Name: ' + name
    output_file = name.replace(' ', '_') + '_report.docx'
    doc_out.add_paragraph(output_str)
    
    total_symptoms = 0
    # Classification starts
    img = Image.open(img_path)

    img = transform_test(img)
    img.unsqueeze_(0)
    img = Variable(img).to(DEVICE)
    out = model_cnn(img)
    # Predict
    _, pred = torch.max(out.data, 1)

    output_result = Refer(classes[pred.data.item()])
    print(classes[pred.data.item()])
    total_symptoms += pred.data.item()
    output_result += '\n'
    # Classification ends
    
    # Detection starts
    results = model_yolo([img_path])

    for result in results: # Should be in a for loop because the 'results' is a list
        im_array = result.plot()
        im = Image.fromarray(im_array[..., ::-1])
        result_json = json.loads(result.tojson())
        output_result += (Refer('并有以下诊断') + ': \n')
        for r in result_json:
            output_result += Refer(r['name'])
            confidence = float(r['confidence']) * 100
            output_result += (Refer('精度') + ': {:.1f}%, '.format(confidence))
            total_symptoms += 1
            if r['name'] == 'OD':
                total_symptoms -= 1
            if confidence < 40:
                output_result += Refer('Lowconf')
            output_result += '\n'
        result.save('tmp.jpg')

    total_symptoms = (3 if total_symptoms > 3 else total_symptoms)

    doc_out.add_paragraph(output_result)
    doc_out.add_paragraph(Refer(severity[total_symptoms]))
    doc_out.add_picture('tmp.jpg', width = Cm(16))

    os.remove('tmp.jpg')
    
    output_result += ('\n' + Refer(severity[total_symptoms]) + '\n')
    output_result += Refer('具体情况已在图中标出')
    # Classification ends
    print(total_symptoms)
    doc_out.save(output_file)
    pythoncom.CoInitialize()
    docx2pdf.convert(output_file, output_file.replace('.docx', '.pdf'))
    pythoncom.CoInitialize()
    return output_result, im # The word output and the image output

iface = gr.Interface(
    fn=predict_image,
    inputs=[
        # gr.Image(type="pil", label="Upload Image")
        gr.Textbox(label="Patient Name"),
        gr.Image(type="filepath", label="Upload Image")
    ],
    outputs=[
        gr.TextArea(label="Word Result"),  # Word output area
        gr.Image(type="pil", label="Image Result") # Image output area
    ],
    css=css, # Still not working now
    title="Pathologic Myopia Detection",
    description="Upload images from the medical eyes scanning. The program will help the doctor to locate the focus and come up with diagnosis"
)

if __name__ == '__main__':
    iface.launch()